import json
from io import BytesIO
from django.http import HttpResponse
from django.utils import timezone
from django.db.models import Avg, Count
from rest_framework import viewsets, permissions, status, generics
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.decorators import action
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from .models import ResearchJob, ResearchReport, AgentStep, MemoryEntry, SourceReliability
from .serializers import (
    ResearchJobSerializer,
    ResearchReportSerializer,
    AgentStepSerializer,
    MemoryEntrySerializer,
    SourceReliabilitySerializer
)
from .tasks import run_research_job
from agent.memory import MemoryManager

# ResearchJob ViewSet
class ResearchJobViewSet(viewsets.ModelViewSet):
    serializer_class = ResearchJobSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return ResearchJob.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        job = serializer.save(user=self.request.user, status='pending')
        # Trigger Celery task asynchronously
        run_research_job.delay(job.id)

    def destroy(self, request, *args, **kwargs):
        job = self.get_object()
        # If running, try to terminate if celery id exists (celery cancel)
        if job.status == 'running' and job.celery_task_id:
            from celery.task.control import revoke
            revoke(job.celery_task_id, terminate=True)
            
        job.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


# ResearchReport ViewSet
class ResearchReportViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = ResearchReportSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return ResearchReport.objects.filter(job__user=self.request.user)

    @action(detail=True, methods=['get'], url_path='export/pdf')
    def export_pdf(self, request, pk=None):
        try:
            report = ResearchReport.objects.get(id=pk, job__user=request.user)
        except ResearchReport.DoesNotExist:
            return Response({"error": "Report not found"}, status=status.HTTP_404_NOT_FOUND)

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="report_{pk}.pdf"'

        doc = SimpleDocTemplate(response, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontSize=22,
            textColor=colors.HexColor('#4f46e5'),
            spaceAfter=15
        )
        h2_style = ParagraphStyle(
            'ReportH2',
            parent=styles['Heading2'],
            fontSize=15,
            textColor=colors.HexColor('#1f2937'),
            spaceBefore=15,
            spaceAfter=6
        )
        body_style = ParagraphStyle(
            'ReportBody',
            parent=styles['BodyText'],
            fontSize=10.5,
            leading=14,
            textColor=colors.HexColor('#374151'),
            spaceAfter=8
        )

        story = []
        story.append(Paragraph(report.title, title_style))
        story.append(Paragraph(f"Confidence Rating: {report.confidence_score:.2f}% | Words: {report.word_count}", body_style))
        story.append(Spacer(1, 15))

        story.append(Paragraph("Executive Summary", h2_style))
        story.append(Paragraph(report.executive_summary, body_style))
        story.append(Spacer(1, 10))

        paragraphs = report.full_report_markdown.split('\n\n')
        for p in paragraphs:
            p = p.strip()
            if not p:
                continue
            if p.startswith('# ') or p.startswith('## ') or p.startswith('### '):
                cleaned = p.replace('#', '').strip()
                story.append(Paragraph(cleaned, h2_style))
            else:
                story.append(Paragraph(p, body_style))

        doc.build(story)
        return response

    @action(detail=True, methods=['get'], url_path='export/docx')
    def export_docx(self, request, pk=None):
        try:
            report = ResearchReport.objects.get(id=pk, job__user=request.user)
        except ResearchReport.DoesNotExist:
            return Response({"error": "Report not found"}, status=status.HTTP_404_NOT_FOUND)

        doc = docx.Document()
        doc.add_heading(report.title, 0)
        
        doc.add_paragraph(f"Confidence Score: {report.confidence_score:.2f}%")
        doc.add_paragraph(f"Word Count: {report.word_count}")
        
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.executive_summary)
        
        paragraphs = report.full_report_markdown.split('\n\n')
        for p in paragraphs:
            p = p.strip()
            if not p:
                continue
            if p.startswith('# ') or p.startswith('## ') or p.startswith('### '):
                cleaned = p.replace('#', '').strip()
                doc.add_heading(cleaned, level=2)
            else:
                doc.add_paragraph(p)
                
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="report_{pk}.docx"'
        return response


# AgentStep List View
class AgentStepListView(generics.ListAPIView):
    serializer_class = AgentStepSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        job_id = self.kwargs['job_id']
        return AgentStep.objects.filter(job_id=job_id, job__user=self.request.user)


# MemoryEntry ViewSet
class MemoryEntryViewSet(viewsets.ModelViewSet):
    serializer_class = MemoryEntrySerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        queryset = MemoryEntry.objects.filter(user=self.request.user)
        query = self.request.query_params.get('query', None)
        if query:
            # Query local Qdrant memory namespace using manager
            memory_mgr = MemoryManager(user_id=str(self.request.user.id))
            hits = memory_mgr.retrieve(query, top_k=10)
            contents = [h["content"] for h in hits]
            queryset = queryset.filter(content__in=contents)
        return queryset

    def perform_create(self, serializer):
        # Allow manual memory entries
        content = serializer.validated_data['content']
        memory_mgr = MemoryManager(user_id=str(self.request.user.id))
        metadata = {
            "source_url": serializer.validated_data.get("source_url", ""),
            "source_tool": "manual",
            "reliability_score": serializer.validated_data.get("reliability_score", 0.5),
            "topic_tags": serializer.validated_data.get("topic_tags", [])
        }
        # Calling store creates both the Qdrant embedding and the Django database object
        point_id = memory_mgr.store(content, metadata)
        
        # Set serializer.instance to the newly created memory entry to return it in the response
        import hashlib
        content_hash = hashlib.sha256(content.encode('utf-8')).hexdigest()
        serializer.instance = MemoryEntry.objects.get(content_hash=content_hash)


# Memory Stats View
class MemoryStatsView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        memories = MemoryEntry.objects.filter(user=request.user)
        total_count = memories.count()
        
        # Calculate top topic tags
        topic_counts = {}
        for mem in memories:
            tags = mem.topic_tags or []
            for tag in tags:
                topic_counts[tag] = topic_counts.get(tag, 0) + 1
        
        sorted_topics = sorted(topic_counts.items(), key=lambda x: x[1], reverse=True)[:5]
        top_topics = [{"topic": k, "count": v} for k, v in sorted_topics]
        
        # Estimated storage (approx 384 dimensions * 4 bytes per float = 1.5 KB per entry)
        storage_kb = total_count * 1.5
        
        return Response({
            "total_count": total_count,
            "top_topics": top_topics,
            "storage_estimate_kb": round(storage_kb, 2)
        })


# Dashboard Stats View
class DashboardStatsView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        user = request.user
        total_jobs = ResearchJob.objects.filter(user=user).count()
        completed_jobs = ResearchJob.objects.filter(user=user, status='done').count()
        total_reports = ResearchReport.objects.filter(job__user=user).count()
        total_memories = MemoryEntry.objects.filter(user=user).count()

        avg_confidence = ResearchReport.objects.filter(job__user=user).aggregate(Avg('confidence_score'))['confidence_score__avg'] or 0.0

        return Response({
            "total_jobs": total_jobs,
            "completed_jobs": completed_jobs,
            "total_reports": total_reports,
            "total_memories": total_memories,
            "avg_confidence_score": round(avg_confidence, 2)
        })
